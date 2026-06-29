"""Tests for DOCXConverter. Fixtures are generated with python-docx at runtime."""

from docx import Document

from pydown.converters.docx import DOCXConverter


def test_headings_and_paragraphs(tmp_path):
    doc = Document()
    doc.add_heading("Title Heading", level=1)
    doc.add_paragraph("A normal paragraph.")
    doc.add_heading("Sub Heading", level=2)
    doc.add_paragraph("Another paragraph.")
    path = tmp_path / "doc.docx"
    doc.save(str(path))

    md = DOCXConverter().convert(str(path))

    assert "# Title Heading" in md
    assert "## Sub Heading" in md
    assert "A normal paragraph." in md
    assert "Another paragraph." in md


def test_document_order_is_preserved(tmp_path):
    doc = Document()
    doc.add_heading("First", level=1)
    doc.add_paragraph("middle text")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "H1"
    table.rows[0].cells[1].text = "H2"
    table.rows[1].cells[0].text = "a"
    table.rows[1].cells[1].text = "b"
    path = tmp_path / "ordered.docx"
    doc.save(str(path))

    md = DOCXConverter().convert(str(path))

    # Heading, then paragraph, then table — in that order.
    assert md.index("# First") < md.index("middle text") < md.index("| H1 | H2 |")


def test_table_renders_as_markdown(tmp_path):
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Name"
    table.rows[0].cells[1].text = "Age"
    table.rows[1].cells[0].text = "Alice"
    table.rows[1].cells[1].text = "30"
    path = tmp_path / "table.docx"
    doc.save(str(path))

    md = DOCXConverter().convert(str(path))

    assert "| Name | Age |" in md
    assert "| --- | --- |" in md
    assert "| Alice | 30 |" in md


def test_empty_document(tmp_path):
    doc = Document()
    path = tmp_path / "empty.docx"
    doc.save(str(path))

    md = DOCXConverter().convert(str(path))

    # No content, no crash — an empty string is acceptable.
    assert md == ""
