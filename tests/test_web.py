"""Tests for the Flask web frontend, driven by Flask's test client.

Uploads are simulated with in-memory files generated at runtime (python-docx),
matching the rest of the suite's "no committed fixtures" approach. ``POST /convert``
returns the Markdown as a downloadable ``.md`` attachment.
"""

import io

import pytest
from docx import Document

from pydown.web.app import create_app


@pytest.fixture
def client():
    app = create_app()
    app.config["TESTING"] = True
    return app.test_client()


def _docx_bytes(heading, paragraph):
    doc = Document()
    doc.add_heading(heading, level=1)
    doc.add_paragraph(paragraph)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def test_index_page_loads(client):
    res = client.get("/")
    assert res.status_code == 200
    body = res.get_data(as_text=True)
    assert "dropzone" in body
    assert "Drag" in body


def test_convert_returns_markdown_as_download(client):
    data = {"file": (_docx_bytes("My Heading", "Body text here"), "sample.docx")}
    res = client.post("/convert", data=data, content_type="multipart/form-data")

    assert res.status_code == 200
    assert res.mimetype == "text/markdown"
    # Served as a downloadable file named after the source.
    assert 'filename="sample.md"' in res.headers["Content-Disposition"]

    markdown = res.get_data(as_text=True)
    assert "# My Heading" in markdown
    assert "Body text here" in markdown


def test_unsupported_extension_returns_400(client):
    data = {"file": (io.BytesIO(b"hello"), "notes.rtf")}
    res = client.post("/convert", data=data, content_type="multipart/form-data")

    assert res.status_code == 400
    assert "Unsupported" in res.get_json()["error"]


def test_missing_file_returns_400(client):
    res = client.post("/convert", data={}, content_type="multipart/form-data")

    assert res.status_code == 400
    assert "error" in res.get_json()
